"""Build the Part B Word and PDF report from precomputed results."""
from __future__ import annotations

import pathlib
from typing import Iterable
from xml.sax.saxutils import escape

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = pathlib.Path(__file__).resolve().parent.parent
REPORT_DIR = ROOT / "report"
DATA_DIR = ROOT / "results" / "data"
TABLE_DIR = ROOT / "results" / "tables"
FIG_DIR = ROOT / "results" / "figures"

FIGURE_CAPTIONS = {
    "growth_of_1_comparison.png": "Figure A1. Growth of $1 across the 12 OOS funds, grouped by universe.",
    "drawdown_comparison.png": "Figure A2. Selected equity and combined drawdown paths; crypto-only funds are excluded.",
    "sharpe_by_fund.png": "Figure A3. Funds are ranked by OOS Sharpe; the dashed line marks the median fund.",
    "combined_risk_parity_weights.png": "Figure A4. Combined Risk Parity crypto sleeve and latest top holdings.",
    "sector_sentiment_index_timeseries.png": "Figure 5. Standalone sector sentiment index over time, 2020-2023.",
    "sector_sentiment_timeseries.png": "Figure A5. Coverage-aware lagged sector sentiment signal, 2020-2023.",
    "fusion_before_after_sharpe.png": "Figure A6. Incremental impact of the sentiment tilt on Sharpe, turnover, and drawdown loss.",
}

EXECUTIVE_SUMMARY = (
    "Headline claim: crypto delivered the strongest terminal growth in the 2021-2023 OOS sample, "
    "but the most suitable retail product is not simply the highest-return fund. SignalHarbor builds "
    "12 funds across equity, crypto, and combined universes, adds a sector-level news sentiment index, "
    "and deploys the results through a Streamlit investor dashboard. Crypto Risk Parity grew $1 to $2.50 "
    "but suffered an -80.4% maximum drawdown, while Equity Equal Weight delivered the highest Sharpe ratio "
    "at 0.82 with a much smaller -20.3% drawdown. Combined Risk Parity is the most balanced combined-universe "
    "candidate in this sample, with a 0.79 Sharpe ratio, -22.9% drawdown, and a modest 8.6% latest crypto sleeve. "
    "The innovation is a custom finance-augmented VADER signal converted into a lagged, coverage-aware sector tilt. "
    "It is transparent and look-ahead safe, but the OOS test slightly reduces Sharpe and raises turnover, so it should "
    "be treated as an analytics layer rather than a production return enhancer."
)

TABLE_SOURCES = {
    "audit": "Source: results/tables/backtest_audit.csv.",
    "metrics": "Source: results/tables/performance_metrics.csv.",
    "fusion": "Source: results/tables/fusion_comparison.csv.",
    "holdings": "Source: results/data/current_holdings.csv.",
}


def pct(x: float, digits: int = 1) -> str:
    return f"{x * 100:.{digits}f}%"


def dec(x: float, digits: int = 2) -> str:
    return f"{x:.{digits}f}"


def money(x: float) -> str:
    return f"${x:.2f}"


def load_tables() -> dict[str, pd.DataFrame]:
    return {
        "metrics": pd.read_csv(TABLE_DIR / "performance_metrics.csv"),
        "audit": pd.read_csv(TABLE_DIR / "backtest_audit.csv"),
        "fusion": pd.read_csv(TABLE_DIR / "fusion_comparison.csv"),
        "sentiment_models": pd.read_csv(TABLE_DIR / "sentiment_model_comparison.csv"),
        "holdings": pd.read_csv(DATA_DIR / "current_holdings.csv"),
    }


def performance_rows(metrics: pd.DataFrame) -> list[list[str]]:
    cols = ["Fund", "Universe", "Method", "Ann. ret.", "Ann. vol.", "Sharpe", "Max DD", "Turnover"]
    rows = [cols]
    for _, r in metrics.sort_values(["universe", "method"]).iterrows():
        rows.append([
            r["fund"],
            r["universe"],
            r["method"],
            pct(r["annualised_return"]),
            pct(r["annualised_volatility"]),
            dec(r["sharpe_ratio"]),
            pct(r["maximum_drawdown"]),
            pct(r["average_turnover"]),
        ])
    return rows


def fusion_rows(fusion: pd.DataFrame) -> list[list[str]]:
    rows = [["Fund", "Experiment", "Ann. ret.", "Vol.", "Sharpe", "Max DD", "Turnover"]]
    for _, r in fusion.iterrows():
        rows.append([
            r["fund"],
            r["experiment"].replace("_", " "),
            pct(r["annualised_return"]),
            pct(r["annualised_volatility"]),
            dec(r["sharpe_ratio"]),
            pct(r["maximum_drawdown"]),
            pct(r["average_turnover"]),
        ])
    return rows


def audit_rows(audit: pd.DataFrame) -> list[list[str]]:
    grouped = audit.groupby("universe").agg(
        funds=("fund_name", "count"),
        rebalances=("rebalance_count", "min"),
        min_oos=("oos_observations", "min"),
        max_oos=("oos_observations", "max"),
        ann=("annualisation_factor", "first"),
        solvers=("solver_failures", "sum"),
        fallbacks=("fallback_count", "sum"),
        lookahead_violations=("strict_no_lookahead_valid", lambda x: int((~x).sum())),
    ).reset_index()
    rows = [["Universe", "Funds", "Rebalances", "OOS daily\nobservations", "Annualisation\nfactor", "Solver\nfailures", "Look-ahead\nviolations"]]
    for _, r in grouped.iterrows():
        obs = str(int(r["min_oos"])) if r["min_oos"] == r["max_oos"] else f"{int(r['min_oos'])}-{int(r['max_oos'])}"
        rows.append([
            r["universe"],
            str(int(r["funds"])),
            str(int(r["rebalances"])),
            obs,
            str(int(r["ann"])),
            str(int(r["solvers"] + r["fallbacks"])),
            str(int(r["lookahead_violations"])),
        ])
    return rows


def design_rows() -> list[list[str]]:
    """Report-level design choices stated in investor-facing language."""
    return [
        ["Design choice", "SignalHarbor setting"],
        ["Universes", "Equity-only, Crypto-only, Combined"],
        ["Methods", "Equal Weight, Minimum Variance, Tangency, Risk Parity"],
        ["Constraint", "Long-only, fully invested: weights are non-negative and sum to one"],
        ["Rebalancing", "Monthly"],
        ["Backtest", "Walk-forward out-of-sample, 2021-2023"],
        ["Estimation", "Expanding window; 2020 initial estimation period; past data only"],
        ["Risk-free rate", "0%"],
        ["Transaction costs", "0% baseline; results are gross of costs and fees"],
        ["Annualisation", "Equity/Combined: 252 trading days; Crypto: 365 calendar days"],
        ["Sentiment use", "Equity-only overlay; one-day-lagged, coverage-adjusted sector signal"],
    ]


def top_holdings_rows(holdings: pd.DataFrame, fund: str, n: int = 8) -> list[list[str]]:
    rows = [["Ticker", "Weight"]]
    h = holdings[holdings["fund"] == fund].sort_values("weight", ascending=False).head(n)
    for _, r in h.iterrows():
        rows.append([r["ticker"], pct(r["weight"], 2)])
    return rows


REPORT_TEXT = [
    (
        "1. Funds and Backtest Design",
        [
            "SignalHarbor is a prototype investment application for retail investors who understand basic finance but do not code. The product offers systematic funds rather than one model portfolio, so users can compare equity-only, crypto-only, and combined equity-plus-crypto exposures before allocating capital.",
            "The model design deliberately separates the three asset universes. Equity-only funds use the 50-stock equity trading calendar, while Crypto-only funds use the seven-day crypto calendar. Combined funds use a stricter trading-product convention. For the Combined universe, crypto returns are first computed on the crypto calendar and then aligned to equity trading dates, so weekend-only crypto returns are not rolled into Monday. Equity and Combined funds use a 252-day annualisation factor, while Crypto funds use 365.",
            "The four construction methods create an interpretable comparison, not just four formulas. Equal Weight is the simple diversification benchmark (DeMiguel et al., 2009). Minimum Variance asks whether covariance information improves capital stability within the mean-variance framework (Markowitz, 1952). Tangency tests whether a mean-return and covariance optimiser can convert high in-sample Sharpe into live OOS performance. Risk Parity asks whether balancing risk contributions helps when assets have very different volatilities (Maillard et al., 2010).",
            "The backtest is a monthly expanding-window out-of-sample design. The initial estimation window is 2020. The first live return date is 2021-01-04 for Equity and Combined funds, and 2021-01-01 for Crypto funds. At each monthly rebalance, weights are formed only from returns observed before the rebalance date; those weights are then held over the next live period. This makes the reported performance an out-of-sample product result rather than an in-sample optimisation fit. Turnover is measured from the drifted pre-rebalance weights to the new target weights. Transaction costs and management fees are assumed to be zero in the baseline results; turnover is reported separately, so the returns should be interpreted as gross rather than net performance.",
        ],
    ),
    (
        "2. Out-of-Sample Results and Fact Sheets",
        [
            "Figure A1 shows that crypto funds delivered the highest terminal growth, but the path was much more volatile than the equity and combined panels. Crypto Risk Parity grew $1 to $2.50 and Crypto Equal Weight grew $1 to $2.42, while the strongest combined fund, Combined Equal Weight, ended at $1.52. The financial interpretation is not simply that crypto is the best asset class. The figure shows that high raw growth is available, but only to investors who can tolerate a much wider path of interim gains and losses. For a retail investor using SignalHarbor, this makes crypto exposure a return engine to size carefully rather than a default core allocation.",
            "Figure A2 adds the missing downside perspective by focusing on selected equity and combined drawdown paths. The selected Tangency funds experienced the deepest losses, with Combined Tangency reaching a -52.7% maximum drawdown and Equity Tangency reaching -38.2%. By comparison, Equity Equal Weight fell -20.3% and Combined Risk Parity fell -22.9%. This supports a practical portfolio decision: the optimiser that uses expected returns may look attractive in-sample, but in this OOS period it created materially worse downside risk. The limitation is that Figure A2 intentionally excludes crypto-only funds so the equity and combined drawdown structure remains readable; crypto risk is evaluated in Table 2 and Figure A1.",
            "Figure A3 converts the return and risk evidence into a cleaner risk-adjusted ranking. Equity Equal Weight has the highest OOS Sharpe at 0.82, followed by Crypto Risk Parity and Combined Risk Parity at 0.79 each, while the median fund is about 0.69. Tangency is weaker across universes: Equity Tangency is 0.65, Crypto Tangency is 0.46, and Combined Tangency is 0.40. This matters because investors should not select the highest terminal-growth product without asking whether the return was earned efficiently. The figure suggests that simple diversification and risk-balancing were more reliable than aggressive mean-return optimisation in this sample.",
            "Table 2 gives the full cross-fund evidence behind these rankings. Crypto Risk Parity has the highest annualised return at 35.8%, but it also has 78.7% volatility and an -80.4% drawdown. Equity Equal Weight is less spectacular, with 12.6% annualised return, 16.2% volatility, and a -20.3% drawdown, but it delivers the best Sharpe ratio. Combined Risk Parity is the most balanced combined-universe candidate, with 12.8% return, 17.2% volatility, a 0.79 Sharpe ratio, and a -22.9% drawdown. The investment implication is that the recommended fund depends on the user's tolerance for path risk, not only on return.",
            "Figure A4 explains why Combined Risk Parity is a plausible multi-asset product rather than simply a crypto bet. Its crypto sleeve stays modest, ranging from about 8.4% to 11.8% of target weight and ending at 8.6% at the latest rebalance. The largest single-name holding is WMT at 3.17%, followed by GILD at 3.10% and MRK at 2.98%, so no individual position dominates the fund. This does not prove the product is safer by itself; that conclusion still depends on Table 2's Sharpe, drawdown, and turnover evidence. It does show that the current holdings structure is interpretable enough for a fact sheet and investor dashboard.",
        ],
    ),
    (
        "3. Sentiment Index",
        [
            "The sentiment claim is deliberately modest: news tone can contain information about investor attention and expectations, but it is a noisy sector signal rather than a standalone return forecast (Tetlock, 2007). The unstructured-data component scores raw equity headlines with two models: baseline VADER and my finance-augmented VADER lexicon (Hutto and Gilbert, 2014). The augmentation adds transparent finance terms such as upgrade, buyback, guidance raise, downgrade, impairment, hawkish, and bankruptcy. This keeps the model explainable and reproducible while adapting a general sentiment tool to financial news language.",
            "The augmented lexicon is treated as a model design choice rather than a guaranteed improvement. I compare it with baseline VADER instead of assuming that more finance words automatically create a better trading signal. In the submitted data, the finance-augmented model changes sector sentiment systematically but incrementally: the average finance-minus-VADER difference is positive in every sector, ranging from +0.24 points in Financials to +0.71 points in Utilities on the 0-100 scale, while sector correlations with baseline VADER remain high. This suggests that the extension changes the interpretation of finance terms without replacing the underlying VADER signal completely.",
            "The pipeline builds evidence in stages. Each headline receives a compound score. Scores are averaged by ticker and date, then sector sentiment is computed as an equal-weight average across active tickers in that sector. This avoids giving a mega-cap with many articles a mechanically larger sentiment weight than a company with fewer headlines. A zero sentiment score may reflect neutral wording or a lexicon miss, while a day with no observed headline is treated as missing rather than neutral. The app displays coverage so users can judge signal reliability.",
            "Figure 5 presents the standalone investor-facing 0-100 sentiment index across the ten equity sectors. It is a readable display of prevailing news tone, with 50 as the neutral midpoint and 21-day averages used only to improve readability. The index is deliberately separate from the trading signal: it helps an investor inspect sector news conditions without implying that the display scale itself is an investment rule. The coverage-aware signal used in the portfolio experiment is shown separately as Figure A5 in the appendix.",
        ],
    ),
    (
        "4. Extension and Innovation: Augmented finVADER and Coverage-Aware Tilt",
        [
            "The main innovation has two connected parts. First, I build a custom finance-augmented VADER model that stays rule-based and auditable rather than using a black-box classifier or simply applying the course example unchanged. Second, I convert that news signal into a portfolio overlay that is lagged, sector-level, and coverage-aware. This makes the extension closer to an implementable investment experiment than a descriptive sentiment chart.",
            "The trading signal is intentionally conservative. For an equity base fund, I first compute the base monthly weights using the same OOS rebalance schedule. I then map each stock to its sector and tilt its base weight using the previous available sector sentiment signal multiplied by the sector's news coverage ratio: Signal(s,t-1) = z(s,t-1) x Coverage(s,t-1). Stock weights are scaled as base_weight x exp(0.15 x Signal) and then re-normalised. Positive lagged sector sentiment increases the sector weight, negative sentiment decreases it, and low coverage automatically weakens the tilt. The exponential form preserves positive weights and keeps the portfolio fully invested after re-normalisation.",
            "The design choices are fixed before evaluation. The tilt strength is set to 0.15 as an illustrative small overlay, not selected after viewing OOS performance. The signal uses a one-day lag and past-only expanding standardisation, so it avoids look-ahead bias. The coverage adjustment also gives the method an economic rationale: a sector with limited headline coverage should not receive the same trading impact as a sector where more tickers have current news.",
            "Table 3 and Figure A6 show that the innovation is not yet an investment improvement. For Equal Weight, the sentiment tilt reduces Sharpe from 0.817 to 0.804 while average monthly turnover rises from 2.7% to 5.6%. For Risk Parity, Sharpe falls from 0.718 to 0.701 while turnover rises from 2.6% to 5.6%. Drawdown protection does not improve: Equal Weight remains about a 20.3% loss, and Risk Parity worsens slightly from 19.6% to 19.7% when drawdown is reported as a positive loss magnitude. The investment conclusion is therefore cautious: the overlay creates a transparent and testable research signal, but its current economic benefit is not strong enough to justify the additional trading burden.",
        ],
    ),
    (
        "5. App and Investor Journey",
        [
            "The Streamlit implementation is a small investment dashboard rather than a code demonstration. Fund Explorer filters the 12 funds by universe and method. Fact Sheet turns each fund into a product page with growth, drawdown, target weights, current holdings, and audit evidence. Portfolio Builder lets the user allocate across funds and view blended historical risk and growth.",
            "Sentiment Analytics makes the news model visible through sector sentiment, coverage, VADER versus finance-augmented VADER, and the base-versus-augmented comparison. Methodology explains expanding-window OOS, calendar differences, sentiment lagging, long-only constraints, and the warning that performance is historical rather than promised.",
            "For deployment, the app reads only precomputed CSV files under results/. It does not download raw data, rerun optimisation, or score headlines at runtime, so the app and report use the same submitted artifacts.",
        ],
    ),
    (
        "6. Critical Reflection and Three Recommendations",
        [
            "First, a conservative investor should start with Equity Equal Weight or Equity Risk Parity rather than crypto-heavy funds. Equity Equal Weight has the highest Sharpe ratio and the equity funds have much shallower drawdowns than crypto-only funds, although the exact ranking may differ in a longer sample.",
            "Second, an investor who wants crypto exposure should consider a Combined fund rather than a standalone Crypto fund. Combined Risk Parity has the strongest Sharpe ratio among the combined funds and a smaller drawdown than crypto-only funds, but its calendar design is conservative: weekend crypto returns are omitted rather than rolled into Monday because the Combined panel keeps only already-computed crypto returns that fall on equity trading dates.",
            "Third, the sentiment overlay should be treated as an analytics feature, not yet as a return-enhancing signal. The current augmented-finVADER tilt is transparent and look-ahead safe, but it raises turnover without improving Sharpe or drawdown. The next version should show which finance terms changed headline scores most often, then test alternative lags, coverage thresholds, lexicon terms, and transaction-cost assumptions before live use.",
        ],
    ),
]


def configure_doc_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_doc_table(doc: Document, rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    for i, value in enumerate(rows[0]):
        hdr[i].text = value
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    if widths:
        for idx, width in enumerate(widths):
            table.columns[idx].width = Inches(width)
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()


def add_doc_source(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(96, 112, 134)


def add_doc_figure(doc: Document, filename: str, width: float = 6.2) -> None:
    doc.add_paragraph(FIGURE_CAPTIONS[filename])
    doc.add_picture(str(FIG_DIR / filename), width=Inches(width))
    doc.add_paragraph()


def build_docx(tables: dict[str, pd.DataFrame]) -> pathlib.Path:
    path = REPORT_DIR / "report.docx"
    doc = Document()
    configure_doc_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("SignalHarbor: Systematic Multi-Asset Funds with News Sentiment")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(23, 50, 77)
    doc.add_paragraph("FINS3645 Part B Report | z5516421")

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(EXECUTIVE_SUMMARY)

    doc.add_paragraph("Backtest audit summary. All funds have 36 monthly rebalances and zero look-ahead violations.")
    add_doc_table(doc, audit_rows(tables["audit"]), widths=[0.8, 0.6, 0.9, 1.25, 1.1, 0.85, 1.0])
    add_doc_source(doc, TABLE_SOURCES["audit"])

    for heading, paragraphs in REPORT_TEXT:
        doc.add_heading(heading, level=1)
        for idx, paragraph in enumerate(paragraphs):
            doc.add_paragraph(paragraph)
            if heading.startswith("1.") and idx == 3:
                doc.add_paragraph("Table 1. Fund Construction and Backtest Design")
                add_doc_table(doc, design_rows(), widths=[1.6, 4.6])
                add_doc_source(doc, "Source: SignalHarbor design settings and results/tables/backtest_audit.csv.")
            if heading.startswith("2.") and idx == 0:
                add_doc_figure(doc, "growth_of_1_comparison.png")
            if heading.startswith("2.") and idx == 1:
                add_doc_figure(doc, "drawdown_comparison.png")
            if heading.startswith("2.") and idx == 2:
                add_doc_figure(doc, "sharpe_by_fund.png")
            if heading.startswith("2.") and idx == 3:
                doc.add_paragraph("Table 2. Performance metrics across all funds, annualised over the OOS period.")
                add_doc_table(doc, performance_rows(tables["metrics"]), widths=[1.6, 0.7, 0.9, 0.65, 0.65, 0.55, 0.6, 0.65])
                add_doc_source(doc, TABLE_SOURCES["metrics"])
            if heading.startswith("2.") and idx == 4:
                add_doc_figure(doc, "combined_risk_parity_weights.png")
            if heading.startswith("3.") and idx == 3:
                add_doc_figure(doc, "sector_sentiment_index_timeseries.png")
            if heading.startswith("4.") and idx == 3:
                doc.add_page_break()
                doc.add_paragraph("Table 3. Base versus sentiment-augmented equity funds.")
                add_doc_table(doc, fusion_rows(tables["fusion"]), widths=[1.75, 1.0, 0.65, 0.6, 0.55, 0.65, 0.65])
                add_doc_source(doc, TABLE_SOURCES["fusion"])
                add_doc_figure(doc, "fusion_before_after_sharpe.png")
        if heading.startswith("2."):
            pass
        if heading.startswith("4."):
            pass

    doc.add_heading("Appendix: Additional Evidence", level=1)

    doc.add_heading("Current Holdings Example", level=2)
    doc.add_paragraph("Table A1. Top current holdings for Combined Risk Parity at the latest rebalance.")
    add_doc_table(doc, top_holdings_rows(tables["holdings"], "Combined Risk Parity"), widths=[2.0, 1.0])
    add_doc_source(doc, TABLE_SOURCES["holdings"])

    doc.add_heading("Coverage-Aware Trading Signal", level=2)
    add_doc_figure(doc, "sector_sentiment_timeseries.png")

    doc.add_heading("References", level=1)
    for ref in [
        "FINS3645 Project Brief, Part B: Funds, Sentiment & App.",
        "Hutto, C. and Gilbert, E. (2014). VADER: A parsimonious rule-based model for sentiment analysis of social media text.",
        "DeMiguel, V., Garlappi, L. and Uppal, R. (2009). Optimal versus naive diversification: How inefficient is the 1/N portfolio strategy?",
        "Maillard, S., Roncalli, T. and Teiletche, J. (2010). The properties of equally weighted risk contribution portfolios.",
        "Markowitz, H. (1952). Portfolio selection.",
        "Tetlock, P. C. (2007). Giving content to investor sentiment: The role of media in the stock market.",
        "Official FINS3645 starter data: equity prices, crypto prices, and news headline files accessed via src/data_access.py.",
    ]:
        doc.add_paragraph(ref)

    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def pdf_styles() -> dict[str, ParagraphStyle]:
    styles = getSampleStyleSheet()
    return {
        "title": ParagraphStyle("TitleCustom", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=18, leading=22, textColor=colors.HexColor("#17324d"), alignment=TA_LEFT, spaceAfter=10),
        "subtitle": ParagraphStyle("SubtitleCustom", parent=styles["Normal"], fontName="Helvetica", fontSize=10, leading=13, textColor=colors.HexColor("#607086"), spaceAfter=12),
        "h1": ParagraphStyle("H1Custom", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=13, leading=16, textColor=colors.HexColor("#2E74B5"), spaceBefore=10, spaceAfter=6),
        "h2": ParagraphStyle("H2Custom", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=11, leading=14, textColor=colors.HexColor("#1F4D78"), spaceBefore=8, spaceAfter=5),
        "body": ParagraphStyle("BodyCustom", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.4, leading=12.2, spaceAfter=6),
        "caption": ParagraphStyle("CaptionCustom", parent=styles["BodyText"], fontName="Helvetica-Oblique", fontSize=8.2, leading=10, textColor=colors.HexColor("#607086"), spaceBefore=4, spaceAfter=5),
    }


def pdf_table(rows: list[list[str]], col_widths: Iterable[float]) -> Table:
    cell_style = ParagraphStyle(
        "TableCell",
        fontName="Helvetica",
        fontSize=6.8,
        leading=8.0,
        alignment=TA_LEFT,
        wordWrap="CJK",
    )
    header_style = ParagraphStyle(
        "TableHeader",
        parent=cell_style,
        fontName="Helvetica-Bold",
        textColor=colors.HexColor("#172033"),
    )
    wrapped = []
    for ridx, row in enumerate(rows):
        style = header_style if ridx == 0 else cell_style
        wrapped.append([Paragraph(escape(str(value)), style) for value in row])
    table = Table(wrapped, colWidths=[w * inch for w in col_widths], repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#172033")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 7.2),
        ("LEADING", (0, 0), (-1, -1), 8.6),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D8DEE9")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    return table


def pdf_image(filename: str, width: float = 6.25) -> Image:
    image = Image(str(FIG_DIR / filename))
    ratio = image.imageHeight / image.imageWidth
    image.drawWidth = width * inch
    image.drawHeight = width * ratio * inch
    return image


def build_pdf(tables: dict[str, pd.DataFrame]) -> pathlib.Path:
    path = REPORT_DIR / "report.pdf"
    styles = pdf_styles()
    story = [
        Paragraph("SignalHarbor: Systematic Multi-Asset Funds with News Sentiment", styles["title"]),
        Paragraph("FINS3645 Part B Report | z5516421", styles["subtitle"]),
        Paragraph("Executive Summary", styles["h1"]),
        Paragraph(EXECUTIVE_SUMMARY, styles["body"]),
        Paragraph("Table 1. Backtest audit summary. All funds have 36 monthly rebalances and zero look-ahead violations.", styles["caption"]),
        pdf_table(audit_rows(tables["audit"]), [0.78, 0.42, 0.7, 0.95, 0.85, 0.8, 0.9]),
        Paragraph(TABLE_SOURCES["audit"], styles["caption"]),
    ]

    for heading, paragraphs in REPORT_TEXT:
        story.append(Paragraph(heading, styles["h1"]))
        for idx, paragraph in enumerate(paragraphs):
            story.append(Paragraph(paragraph, styles["body"]))
            if heading.startswith("2.") and idx == 0:
                story.append(Paragraph(FIGURE_CAPTIONS["growth_of_1_comparison.png"], styles["caption"]))
                story.append(pdf_image("growth_of_1_comparison.png"))
                story.append(Spacer(1, 0.08 * inch))
            if heading.startswith("2.") and idx == 1:
                story.append(Paragraph(FIGURE_CAPTIONS["drawdown_comparison.png"], styles["caption"]))
                story.append(pdf_image("drawdown_comparison.png"))
                story.append(Spacer(1, 0.08 * inch))
            if heading.startswith("2.") and idx == 2:
                story.append(Paragraph(FIGURE_CAPTIONS["sharpe_by_fund.png"], styles["caption"]))
                story.append(pdf_image("sharpe_by_fund.png"))
                story.append(Spacer(1, 0.08 * inch))
            if heading.startswith("2.") and idx == 3:
                story.append(Paragraph("Table 2. Performance metrics across all funds, annualised over the OOS period.", styles["caption"]))
                story.append(pdf_table(performance_rows(tables["metrics"]), [1.34, 0.55, 0.78, 0.54, 0.54, 0.46, 0.5, 0.55]))
                story.append(Paragraph(TABLE_SOURCES["metrics"], styles["caption"]))
            if heading.startswith("2.") and idx == 4:
                story.append(Paragraph(FIGURE_CAPTIONS["combined_risk_parity_weights.png"], styles["caption"]))
                story.append(pdf_image("combined_risk_parity_weights.png"))
                story.append(Spacer(1, 0.08 * inch))
            if heading.startswith("3.") and idx == 3:
                story.append(Paragraph(FIGURE_CAPTIONS["sector_sentiment_index_timeseries.png"], styles["caption"]))
                story.append(pdf_image("sector_sentiment_index_timeseries.png"))
                story.append(Spacer(1, 0.08 * inch))
            if heading.startswith("4.") and idx == 3:
                story.append(PageBreak())
                story.append(Paragraph("Table 3. Base versus sentiment-augmented equity funds.", styles["caption"]))
                story.append(pdf_table(fusion_rows(tables["fusion"]), [1.75, 0.85, 0.55, 0.5, 0.5, 0.55, 0.55]))
                story.append(Paragraph(TABLE_SOURCES["fusion"], styles["caption"]))
                story.append(Paragraph(FIGURE_CAPTIONS["fusion_before_after_sharpe.png"], styles["caption"]))
                story.append(pdf_image("fusion_before_after_sharpe.png"))
                story.append(Spacer(1, 0.08 * inch))

    story.append(Paragraph("Appendix: Additional Evidence", styles["h1"]))
    story.append(Paragraph("Current Holdings Example", styles["h2"]))
    story.append(Paragraph("Table A1. Top current holdings for Combined Risk Parity at the latest rebalance.", styles["caption"]))
    story.append(pdf_table(top_holdings_rows(tables["holdings"], "Combined Risk Parity"), [1.5, 1.0]))
    story.append(Paragraph(TABLE_SOURCES["holdings"], styles["caption"]))
    story.append(Paragraph("Coverage-Aware Trading Signal", styles["h2"]))
    story.append(Paragraph(FIGURE_CAPTIONS["sector_sentiment_timeseries.png"], styles["caption"]))
    story.append(pdf_image("sector_sentiment_timeseries.png"))

    story.append(Paragraph("References", styles["h1"]))
    for ref in [
        "FINS3645 Project Brief, Part B: Funds, Sentiment & App.",
        "Hutto, C. and Gilbert, E. (2014). VADER: A parsimonious rule-based model for sentiment analysis of social media text.",
        "DeMiguel, V., Garlappi, L. and Uppal, R. (2009). Optimal versus naive diversification: How inefficient is the 1/N portfolio strategy?",
        "Maillard, S., Roncalli, T. and Teiletche, J. (2010). The properties of equally weighted risk contribution portfolios.",
        "Markowitz, H. (1952). Portfolio selection.",
        "Tetlock, P. C. (2007). Giving content to investor sentiment: The role of media in the stock market.",
        "Official FINS3645 starter data: equity prices, crypto prices, and news headline files accessed via src/data_access.py.",
    ]:
        story.append(Paragraph(ref, styles["body"]))

    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.72 * inch,
        bottomMargin=0.72 * inch,
        title="SignalHarbor FINS3645 Part B Report",
    )
    doc.build(story)
    return path


def main() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    tables = load_tables()
    docx_path = build_docx(tables)
    pdf_path = build_pdf(tables)
    print(f"saved: {docx_path.relative_to(ROOT)}")
    print(f"saved: {pdf_path.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
